"""
docx_parser.py — Document loading, table mapping, numeric detection.
"""
import re
import logging
from docx import Document

logger = logging.getLogger(__name__)

# Matches cells that are purely numeric/financial (not narrative text)
_NUMERIC_RE = re.compile(
    r"""^[\s$%,()\d\-—–.]*$""",
    re.VERBOSE,
)
# Must contain at least one digit or be an explicit zero-stand-in
_ZERO_STAND_INS = {"—", "–", "-", ""}


def is_numeric_cell(text: str) -> bool:
    """
    Returns True if the cell contains only financial numerics,
    currency symbols, punctuation, or zero stand-ins.
    """
    stripped = text.strip()
    if stripped in _ZERO_STAND_INS:
        return len(stripped) > 0  # only "—" "–" count, not truly empty
    return bool(_NUMERIC_RE.match(stripped)) and any(c.isdigit() for c in stripped)


def load_document(path: str) -> Document:
    """Open and return a python-docx Document object."""
    return Document(path)


def get_unique_cells_in_row(row) -> list:
    """
    Return cells in a row, deduplicating merged cells by their _tc XML element identity.
    Prevents double-processing when a cell spans multiple columns.
    """
    seen = set()
    unique = []
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id not in seen:
            seen.add(tc_id)
            unique.append(cell)
    return unique


def _cell_text(cell) -> str:
    return cell.text.strip()


def classify_table_columns(table, source_config: dict, target_config: dict) -> list:
    """
    Inspect header rows of a table and return a list of dicts:
        [{col_index: int, role: str}, ...]

    Roles:
        LABEL           — row label column (col 0 typically)
        CURRENT_PERIOD  — target period (e.g. March 31, 2026)
        COMPARABLE      — comparable prior-year period (e.g. March 31, 2025)
        YTD_CURRENT     — 6-month YTD current
        YTD_COMPARABLE  — 6-month YTD comparable
        PRIOR_YEAR_END  — balance sheet prior year end (Sep 30, 2025)
        SEPARATOR       — empty separator column
        UNKNOWN         — could not classify

    Uses table.cell(r, c) directly to correctly handle merged cells that span
    multiple column indices (common in financial statement headers).
    """
    src_period = source_config["period_label"]
    src_comparable = source_config["comparable_label"]
    src_prior_ye = source_config["prior_year_end_label"]
    tgt_period = target_config["period_label"]
    tgt_comparable = target_config["comparable_label"]
    tgt_prior_ye = target_config["prior_year_end_label"]

    num_cols = len(table.columns)
    num_header_rows = min(3, len(table.rows))
    col_roles = []

    for col_idx in range(num_cols):
        # Use table.cell() to correctly retrieve text for merged cells
        texts = []
        for row_idx in range(num_header_rows):
            try:
                texts.append(table.cell(row_idx, col_idx).text.strip().lower())
            except Exception:
                texts.append("")
        combined = " ".join(texts)

        if col_idx == 0:
            role = "LABEL"
        elif not combined.strip():
            role = "SEPARATOR"
        elif src_prior_ye.lower() in combined or tgt_prior_ye.lower() in combined or "september 30" in combined:
            role = "PRIOR_YEAR_END"
        elif src_period.lower() in combined or tgt_period.lower() in combined:
            role = "CURRENT_PERIOD"
        elif src_comparable.lower() in combined or tgt_comparable.lower() in combined:
            role = "COMPARABLE"
        elif "six months" in combined or "6 months" in combined:
            if "2026" in combined or "2025" in combined:
                role = "YTD_CURRENT" if "2026" in combined else "YTD_COMPARABLE"
            else:
                role = "UNKNOWN"
        else:
            role = "UNKNOWN"

        col_roles.append({"col_index": col_idx, "role": role})

    return col_roles


def extract_table_map(doc: Document) -> list:
    """
    Diagnostic: returns metadata for every table in the document.
    Returns list of dicts with index, rows, cols, and first few header cell texts.
    """
    result = []
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        headers = []
        for r_idx in range(min(3, rows)):
            row_cells = get_unique_cells_in_row(table.rows[r_idx])
            headers.append([c.text.strip()[:60] for c in row_cells])
        result.append({
            "index": i,
            "rows": rows,
            "cols": cols,
            "headers": headers,
        })
    return result


# ---------------------------------------------------------------------------
# Keyword signatures for company-agnostic table detection
# ---------------------------------------------------------------------------
_TABLE_SIGNATURES = {
    "balance_sheet": {
        "keywords": ["total assets", "total liabilities", "stockholders' equity"],
        "min_matches": 3,
    },
    "income_statement": {
        "keywords": ["revenues", "net income", "earnings per share"],
        "min_matches": 2,
    },
    "cash_flow": {
        "keywords": ["operating activities", "investing activities", "financing activities"],
        "min_matches": 3,
    },
    "stockholders_equity": {
        "keywords": ["common stock", "retained earnings", "paid-in capital"],
        "min_matches": 3,
    },
    "eps": {
        "keywords": ["earnings per share", "weighted average"],
        "min_matches": 2,
    },
}


def _collect_scan_texts(table) -> list:
    """
    Collect lowercased text from:
      - All cells in the first 5 rows
      - The first cell of every row (label column)
    Returns a flat list of strings.
    """
    texts = []
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # All cells in the first 5 rows
    for r_idx in range(min(5, num_rows)):
        for c_idx in range(num_cols):
            try:
                texts.append(table.cell(r_idx, c_idx).text.strip().lower())
            except Exception:
                pass

    # First column of every row beyond row 5
    for r_idx in range(5, num_rows):
        try:
            texts.append(table.cell(r_idx, 0).text.strip().lower())
        except Exception:
            pass

    return texts


def detect_financial_tables(doc: Document) -> dict:
    """
    Scan all tables in *doc* for keyword signatures and return a dict mapping
    table type names to their 0-based table indices.

    Return structure
    ----------------
    {
        "balance_sheet":       <int>,   # present only if detected
        "income_statement":    <int>,
        "cash_flow":           <int>,
        "stockholders_equity": <int>,
        "eps":                 <int>,
        "detected_at": {
            <table_type>: [<matched_keyword>, ...],
            ...
        },
    }

    Detection strategy
    ------------------
    For each table, scan all cells in the first 5 rows plus the entire first
    column.  Check how many of the type's keywords appear in that text corpus.
    If the count meets the minimum threshold the table is matched to that type.
    The first matching table wins for each type.

    Special rule for "eps": a table that already matched "income_statement"
    cannot also match "eps" — the EPS note table is a secondary table.
    """
    result = {}
    detected_at = {}

    for tbl_idx, table in enumerate(doc.tables):
        scan_texts = _collect_scan_texts(table)
        combined = " ".join(scan_texts)

        for table_type, sig in _TABLE_SIGNATURES.items():
            # Skip if this type is already matched
            if table_type in result:
                continue

            # "eps" must not be the same table already claimed as income_statement
            if table_type == "eps" and result.get("income_statement") == tbl_idx:
                continue

            matched = [kw for kw in sig["keywords"] if kw in combined]
            if len(matched) >= sig["min_matches"]:
                result[table_type] = tbl_idx
                detected_at[table_type] = matched
                logger.debug(
                    "detect_financial_tables: table %d → %s (matched: %s)",
                    tbl_idx, table_type, matched,
                )

    # Warn for any type that was not found
    for table_type in _TABLE_SIGNATURES:
        if table_type not in result:
            logger.warning(
                "detect_financial_tables: could not detect '%s' table by keywords", table_type
            )

    result["detected_at"] = detected_at
    return result
