"""
text_updater.py — Paragraph/run-level date & period text replacement.
"""
import logging
from docx.enum.text import WD_COLOR_INDEX

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Month words used in "N Months Ended" headers
# ---------------------------------------------------------------------------
_MONTHS_WORD = {3: "Three", 6: "Six", 9: "Nine", 12: "Twelve"}


def build_rules(source_config: dict, target_config: dict) -> list:
    """
    Build an ordered list of (old, new) text replacement rules derived from
    the source and target period configurations.  More-specific strings come
    first so a longer phrase is matched before its shorter sub-strings.

    Handles:
    - "N Months Ended <date>" headers in three case variants
    - "Year Ended" / "Fiscal Year Ended" language when ytd_months == 12
    - Plain date labels in normal and ALL-CAPS variants
    - Quarter ordinal references (first/second/…) with several case forms
    - Filing date → placeholder
    """
    src_p  = source_config["period_label"]        # e.g. "December 31, 2025"
    src_c  = source_config["comparable_label"]    # e.g. "December 31, 2024"
    src_q  = source_config["quarter_name"].lower()  # e.g. "first"
    src_fd = source_config.get("filing_date", "")
    src_mo = source_config.get("ytd_months", 3)

    tgt_p  = target_config["period_label"]
    tgt_c  = target_config["comparable_label"]
    tgt_q  = target_config["quarter_name"].lower()
    tgt_fd = target_config.get("filing_date", "[FILING DATE]")
    tgt_mo = target_config.get("ytd_months", src_mo)

    src_mo_word = _MONTHS_WORD.get(src_mo, f"{src_mo}")
    tgt_mo_word = _MONTHS_WORD.get(tgt_mo, f"{tgt_mo}")

    rules = []

    # ── Annual (12-month) "Year Ended" language ──────────────────────────────
    # Inject these before "N Months Ended" so the more specific phrases match
    # first.  Active whenever source ytd_months == 12.
    if src_mo == 12:
        rules += _build_year_ended_rules(src_p, tgt_p, src_c, tgt_c)

    # ── "N Months Ended <date>" — 3 case variants × 2 periods × 2 months words ──
    # We cover both the single-quarter header ("Three Months Ended") and any
    # YTD header that appears in the same doc ("Six Months Ended" for Q2, etc.)
    for src_mw, tgt_mw in [(src_mo_word, tgt_mo_word), ("Three", "Three")]:
        # The second pair keeps the 3-month header intact when src==tgt (no-op)
        for src_date, tgt_date in [(src_p, tgt_p), (src_c, tgt_c)]:
            rules += [
                (f"{src_mw} Months Ended {src_date}",         f"{tgt_mw} Months Ended {tgt_date}"),
                (f"{src_mw.lower()} months ended {src_date}", f"{tgt_mw.lower()} months ended {tgt_date}"),
                (f"{src_mw.upper()} MONTHS ENDED {src_date.upper()}", f"{tgt_mw.upper()} MONTHS ENDED {tgt_date.upper()}"),
            ]

    # Deduplicate while preserving order (dict trick)
    rules = list(dict.fromkeys(rules))

    # ── Plain date labels ──
    for src_date, tgt_date in [(src_p, tgt_p), (src_c, tgt_c)]:
        rules += [
            (src_date,              tgt_date),
            (src_date.upper(),      tgt_date.upper()),
        ]

    # ── Quarter ordinal references ──
    if src_q != tgt_q:
        # Extract fiscal year numbers that appear in the document
        # e.g. "first quarter of fiscal 2026" → "second quarter of fiscal 2026"
        src_years = _extract_years(src_p, src_c)
        for yr in src_years:
            rules += [
                (f"{src_q} quarter of fiscal {yr}",         f"{tgt_q} quarter of fiscal {yr}"),
                (f"{src_q.title()} quarter of fiscal {yr}", f"{tgt_q.title()} quarter of fiscal {yr}"),
                (f"{src_q.title()} Quarter of Fiscal {yr}", f"{tgt_q.title()} Quarter of Fiscal {yr}"),
            ]
        # Generic quarter-only (no year) — must come after year-specific
        rules += [
            (f"{src_q} quarter",         f"{tgt_q} quarter"),
            (f"{src_q.title()} Quarter", f"{tgt_q.title()} Quarter"),
            (f"{src_q.upper()} QUARTER", f"{tgt_q.upper()} QUARTER"),
        ]

    # ── Filing date ──
    if src_fd and src_fd != tgt_fd:
        rules.append((src_fd, tgt_fd))

    # Final dedup
    return list(dict.fromkeys(rules))


def _build_year_ended_rules(src_p: str, tgt_p: str, src_c: str, tgt_c: str) -> list:
    """
    Return replacement rules for "Year Ended" and "Fiscal Year Ended" language.
    Covers common capitalisation variants and "For the year ended" constructions.
    Called automatically by build_rules() when ytd_months == 12, and also
    used directly by build_annual_rules().
    """
    rules = []
    for src_date, tgt_date in [(src_p, tgt_p), (src_c, tgt_c)]:
        src_upper = src_date.upper()
        tgt_upper = tgt_date.upper()
        rules += [
            # "Fiscal Year Ended" variants — check before plain "Year Ended"
            (f"Fiscal Year Ended {src_date}",         f"Fiscal Year Ended {tgt_date}"),
            (f"fiscal year ended {src_date}",         f"fiscal year ended {tgt_date}"),
            (f"FISCAL YEAR ENDED {src_upper}",        f"FISCAL YEAR ENDED {tgt_upper}"),
            # "For the year ended" variants
            (f"For the year ended {src_date}",        f"For the year ended {tgt_date}"),
            (f"for the year ended {src_date}",        f"for the year ended {tgt_date}"),
            (f"FOR THE YEAR ENDED {src_upper}",       f"FOR THE YEAR ENDED {tgt_upper}"),
            # Plain "Year Ended" variants
            (f"Year Ended {src_date}",                f"Year Ended {tgt_date}"),
            (f"year ended {src_date}",                f"year ended {tgt_date}"),
            (f"YEAR ENDED {src_upper}",               f"YEAR ENDED {tgt_upper}"),
        ]
    return rules


def build_annual_rules(source_config: dict, target_config: dict) -> list:
    """
    Build replacement rules for a 10-K annual roll-forward.

    Handles "Year Ended" and "Fiscal Year Ended" language in addition to all
    the standard date-label replacements produced by build_rules().  The
    "N Months Ended" patterns are omitted because they do not appear in 10-K
    income-statement headers.

    Returns the same format as build_rules(): a list of (old, new) tuples.
    """
    src_p  = source_config["period_label"]
    src_c  = source_config["comparable_label"]
    src_fd = source_config.get("filing_date", "")

    tgt_p  = target_config["period_label"]
    tgt_c  = target_config["comparable_label"]
    tgt_fd = target_config.get("filing_date", "[FILING DATE]")

    rules = []

    # ── "Year Ended" / "Fiscal Year Ended" / "For the year ended" ────────────
    rules += _build_year_ended_rules(src_p, tgt_p, src_c, tgt_c)

    # Deduplicate (preserves order)
    rules = list(dict.fromkeys(rules))

    # ── Plain date labels ──
    for src_date, tgt_date in [(src_p, tgt_p), (src_c, tgt_c)]:
        rules += [
            (src_date,         tgt_date),
            (src_date.upper(), tgt_date.upper()),
        ]

    # ── Filing date ──
    if src_fd and src_fd != tgt_fd:
        rules.append((src_fd, tgt_fd))

    return list(dict.fromkeys(rules))


def _extract_years(period_label: str, comparable_label: str) -> list:
    """Pull 4-digit year strings from the two period labels."""
    import re
    years = re.findall(r'\b(20\d{2})\b', period_label + " " + comparable_label)
    return list(dict.fromkeys(years))  # unique, preserving order


# ---------------------------------------------------------------------------
# Legacy constant — kept for any callers that import it directly.
# engine.py now calls build_rules() instead.
# ---------------------------------------------------------------------------
DEFAULT_RULES = build_rules(
    source_config={
        "period_label": "December 31, 2025", "comparable_label": "December 31, 2024",
        "quarter_name": "first", "filing_date": "February 17, 2026", "ytd_months": 3,
    },
    target_config={
        "period_label": "March 31, 2026", "comparable_label": "March 31, 2025",
        "quarter_name": "second", "filing_date": "[FILING DATE]", "ytd_months": 6,
    },
)


def replace_in_run_text(run, rules: list) -> int:
    """Apply replacement rules to a single Run. Returns number of replacements made."""
    count = 0
    for old, new in rules:
        if old in run.text:
            run.text = run.text.replace(old, new)
            count += 1
    return count


def _full_paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def consolidate_runs(paragraph):
    """
    Merge all runs into runs[0], preserving runs[0] formatting.
    Used when a target string spans multiple runs.
    Only call after confirming a cross-run match exists.
    """
    if len(paragraph.runs) <= 1:
        return
    full_text = _full_paragraph_text(paragraph)
    paragraph.runs[0].text = full_text
    # Clear subsequent runs
    for run in paragraph.runs[1:]:
        run.text = ""


def update_paragraph_text(paragraph, rules: list) -> int:
    """
    Apply rules run-by-run first; if a target string spans runs,
    consolidate then re-apply. Returns total replacement count.
    """
    count = 0
    for run in paragraph.runs:
        count += replace_in_run_text(run, rules)

    # Check if any rule target STILL spans run boundaries after individual pass.
    # This handles cases where one rule matched individually but another target
    # still spans multiple runs (e.g. same paragraph has two date references,
    # one in a single run and one fragmented across runs).
    full_text = _full_paragraph_text(paragraph)
    still_has_cross_run = any(old in full_text for old, _ in rules)
    if still_has_cross_run:
        consolidate_runs(paragraph)
        extra = 0
        for run in paragraph.runs:
            extra += replace_in_run_text(run, rules)
        count += extra

    return count


class _XmlRunAdapter:
    """Duck-type adapter giving update_paragraph_text a run-like object from a w:r XML element."""
    def __init__(self, r_elem):
        from docx.oxml.ns import qn
        self._r   = r_elem
        self._qn  = qn
        self._XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

    @property
    def text(self) -> str:
        return "".join(t.text or "" for t in self._r.iter(self._qn('w:t')))

    @text.setter
    def text(self, value: str):
        t_elems = list(self._r.iter(self._qn('w:t')))
        if not t_elems:
            return
        t_elems[0].text = value
        for t in t_elems[1:]:
            t.text = ""
        # Preserve whitespace at string boundaries
        if value and (value[0] == ' ' or value[-1] == ' '):
            t_elems[0].set(self._XML_SPACE, 'preserve')


class _XmlParagraphAdapter:
    """Duck-type adapter giving update_paragraph_text a paragraph-like object from a w:p element."""
    def __init__(self, p_elem):
        from docx.oxml.ns import qn
        self.runs = [_XmlRunAdapter(r) for r in p_elem.iter(qn('w:r'))]


def _get_notes_paragraphs(doc) -> list:
    """
    Return _XmlParagraphAdapter objects for every w:p in footnotes.xml
    and endnotes.xml (both parts are optional; silently skipped if absent).
    """
    from docx.oxml.ns import qn
    _FOOTNOTES = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
    _ENDNOTES  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes'
    paras = []
    try:
        for rel in doc.part.rels.values():
            if rel.reltype in (_FOOTNOTES, _ENDNOTES):
                for p_elem in rel.target_part._element.iter(qn('w:p')):
                    paras.append(_XmlParagraphAdapter(p_elem))
    except Exception as exc:
        logger.warning("Could not access footnotes/endnotes: %s", exc)
    return paras


def update_all_paragraphs(doc, rules: list = None) -> int:
    """
    Apply text replacement to all paragraphs in the document,
    including paragraphs inside table cells (which doc.paragraphs skips).
    Returns total replacement count.
    """
    if rules is None:
        rules = DEFAULT_RULES

    total = 0

    # Body paragraphs
    for para in doc.paragraphs:
        total += update_paragraph_text(para, rules)

    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += update_paragraph_text(para, rules)

    logger.info("Text replacement pass: %d substitutions made", total)

    # Footnote / endnote paragraphs (word/footnotes.xml, word/endnotes.xml)
    note_paras = _get_notes_paragraphs(doc)
    for para in note_paras:
        total += update_paragraph_text(para, rules)
    if note_paras:
        logger.info("Text replacement pass (footnotes/endnotes): %d paragraphs", len(note_paras))

    return total


def highlight_financial_paragraphs_for_review(doc) -> int:
    """
    Apply yellow highlight to runs containing '$' in non-table body paragraphs
    (MD&A narrative). Returns count of highlighted runs.
    """
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if "$" in run.text:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                count += 1
    logger.info("Highlighted %d MD&A runs containing '$' for manual review", count)
    return count
